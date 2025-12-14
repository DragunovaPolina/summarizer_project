import tkinter as tk
from tkinter import filedialog, messagebox
import os

import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

import docx
from fpdf import FPDF
import PyPDF2

# ________НАСТРОЙКА МОДЕЛИ________

MODEL_NAME = "cointegrated/rut5-base-absum"

tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)

# ________ЧТЕНИЕ ФАЙЛОВ________

def read_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def read_docx(path):
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def read_pdf(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text

def read_file(path):
    if path.endswith(".txt"):
        return read_txt(path)
    elif path.endswith(".docx"):
        return read_docx(path)
    elif path.endswith(".pdf"):
        return read_pdf(path)
    else:
        raise ValueError("Неподдерживаемый формат файла")


# ________РЕЗЮМИРОВАНИЕ________

def summarize_text(text, max_len=300):
    inputs = tokenizer(
        text,
        return_tensors="pt",
        truncation=True,
        max_length=1024
    )

    with torch.no_grad():
        summary_ids = model.generate(
            inputs["input_ids"],
            max_length=max_len,
            min_length=150,
            num_beams=6,
            length_penalty=2.0,
            early_stopping=True,
            no_repeat_ngram_size=3
        )

    return tokenizer.decode(summary_ids[0], skip_special_tokens=True)


# ________GUI________

class SummarizerApp:

    def __init__(self, root):
        self.root = root
        self.root.title("Резюмирование текста")
        self.root.geometry("850x600")

        self.file_path = None

        tk.Label(root, text="Программа автоматического резюмирования текста", font=("Arial", 15, "bold")).pack(pady=10)

        tk.Button(root, text="Выбрать файл (TXT / DOCX / PDF)", command=self.select_file).pack()
        self.file_label = tk.Label(root, text="Файл не выбран")
        self.file_label.pack(pady=5)

        self.text_area = tk.Text(root, wrap="word", height=25)
        self.text_area.pack(padx=10, pady=10, fill="both", expand=True)

        tk.Button(root, text="Резюмировать", command=self.run_summarization).pack(pady=5)
        tk.Button(root, text="Сохранить результат", command=self.save_result).pack(pady=5)

    # ----------------------------------------------
    def select_file(self):
        self.file_path = filedialog.askopenfilename(filetypes=[("Документы", "*.txt *.docx *.pdf")])
        if self.file_path:
            self.file_label.config(text=os.path.basename(self.file_path))

    # ----------------------------------------------
    def run_summarization(self):
        if not self.file_path:
            messagebox.showwarning("Ошибка", "Сначала выберите файл")
            return

        try:
            text = read_file(self.file_path)
            summary = summarize_text(text)
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(tk.END, summary)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    # ----------------------------------------------
    def save_result(self):
        summary = self.text_area.get(1.0, tk.END).strip()
        if not summary:
            messagebox.showwarning("Ошибка", "Нет текста для сохранения")
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text file", "*.txt"), ("Word Document", "*.docx"), ("PDF file", "*.pdf")]
        )

        if path:
            ext = os.path.splitext(path)[1].lower()
            try:
                if ext == ".txt":
                    with open(path, "w", encoding="utf-8") as f:
                        f.write(summary)
                elif ext == ".docx":
                    doc = docx.Document()
                    doc.add_paragraph(summary)
                    doc.save(path)
                elif ext == ".pdf":
                    # Создание PDF с поддержкой кириллицы
                    pdf = FPDF()
                    pdf.add_page()
                
                    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
                    pdf.set_font("DejaVu", size=12)
                    for line in summary.split("\n"):
                        pdf.multi_cell(0, 6, line)
                    pdf.output(path)
                messagebox.showinfo("Готово", "Файл успешно сохранён")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))



# ________ЗАПУСК________

if __name__ == "__main__":
    root = tk.Tk()
    app = SummarizerApp(root)
    root.mainloop()